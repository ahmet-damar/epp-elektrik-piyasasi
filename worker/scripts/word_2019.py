"""EPP — 2019 Word (.docx) EPDK aylık raporları: TEK SEFERLİK tarihsel
aktarım tarifi. Bkz. dokumanlar/08_word_2016_2022_kapsam.md (teşhis) ve
worker/scripts/word_2020.py (yapısal temel — bu dosya onun kopyası, YIL
BAZLI AYRI TARİF ilkesi bilinçli tercih).

**Ortam notu (2026-09-04):** Bu makinede pandas'ın derlenmiş bir bileşeni
Windows Akıllı Uygulama Denetimi (Değerlendirme modu) tarafından
engellenmişti (`worker.parser` çalışmıyordu) — Miniconda ile (`C:\\Users\\
adama\\miniconda3\\envs\\epp`, conda-forge kanalı) çözüldü, imzalı
derlemeler engellenmiyor. Tüm dry-run/test/yükleme komutları bu ortamın
Python'ıyla çalıştırılmalı: `C:\\Users\\adama\\miniconda3\\envs\\epp\\
python.exe`.

**GECE-BOYU GÖZETİMSİZ ÇALIŞMA KURALLARI (Ahmet'in talimatı):**
- `pipeline.batch_onayla()` / `worker/scripts/onayla.py` BU DOSYADA HİÇ
  ÇAĞRILMAZ.
- Taksonomi kararı BAŞTAN dahil (RENAME, 2021/2022'de verildi).
- Ay/yıl doğrulaması kapak paragrafından (word_2020.py ile AYNI — T11
  başlığı "Döneminde" içermiyor).
- T10 GERÇEKTEN il-ONLY görünüyor (Ocak+Aralık örnekleri kontrol edildi,
  dokumanlar/08) — `isle_ay()` T11'i T10'dan BAĞIMSIZ okuyor.
- **YENİ bulgu (2018/2019'a özgü): T4 tablosunda "Güneş" TEK kolon DEĞİL**
  — "Güneş (Fotovoltaik)" VE "Güneş (Yoğunlş.)" AYRI iki kolon, ikisi de
  AYNI kanonik "Güneş" kaynağına toplanmalı (worker/parser.py'nin Excel
  tarafındaki Akarsu+Barajlı→Hidrolik birleştirmesiyle AYNI ilke,
  test_parser.py'de örneklenmiş). `t4_oku()` bu yüzden kolon bazlı değil,
  İL BAŞINA KAYNAK TOPLAMI mantığıyla yazıldı (aynı kaynağa eşlenen birden
  fazla kolon TOPLANIR, ayrı satır ÜRETİLMEZ — fact_uretim'in doğal
  anahtarında {il,tarih,kaynak,lisans} tekil olmalı).

Kullanım:
    python -m worker.scripts.word_2019 --ay 6                  # tek ay, gerçek yükleme
    python -m worker.scripts.word_2019                          # 2019'un TÜMÜ
    python -m worker.scripts.word_2019 --dry-run                # yalnız parse, DB'ye YAZMA
    python -m worker.scripts.word_2019 --t4                     # T11/T10 yerine T4
"""

from __future__ import annotations

import argparse
import re
import sys
from io import BytesIO
from pathlib import Path

import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document

from worker import ingest, kpi, pipeline
from worker.db import get_database_url  # import yan etkisi: .env yüklenir
from worker.parser import grup_esle as _excel_grup_esle
from worker.parser import il_kodu_bul, kaynak_esle, parse_sayi
from worker.scripts.word_ortak import (
    basliklari_topla,
    genel_toplam_satirini_oku,
    grup_kolonlarini_coz,
    hedef_donem_kolonu_bul,
    t4_tablosunu_bul,
    tek_aday_bul,
)

KLASOR_VARSAYILAN = Path(r"C:\Users\adama\Downloads\EPDK Verileri")

# ay(int) -> dosya adı. dokumanlar/08_word_2016_2022_kapsam.md'deki envanter
# taramasıyla bulundu. Ekim dosyası "_Content_Media_" öneki taşıyor (diğer
# tüm aylar "_Content_FastAccess_") — kozmetik, dosya normal açılıyor.
MANIFEST_2019: dict[int, str] = {
    1: "_PortalAdmin_Uploads_Content_FastAccess_60a870be45976.docx",
    2: "_PortalAdmin_Uploads_Content_FastAccess_90f80bc165769.docx",
    3: "_PortalAdmin_Uploads_Content_FastAccess_c3c0bb6c71292.docx",
    4: "_PortalAdmin_Uploads_Content_FastAccess_7671686a66292.docx",
    5: "_PortalAdmin_Uploads_Content_FastAccess_8ac4639a15500.docx",
    6: "_PortalAdmin_Uploads_Content_FastAccess_7356f49d47309.docx",
    7: "_PortalAdmin_Uploads_Content_FastAccess_4429791882972.docx",
    8: "_PortalAdmin_Uploads_Content_FastAccess_a5e2015f46581.docx",
    9: "_PortalAdmin_Uploads_Content_FastAccess_a5c7891379595.docx",
    10: "_PortalAdmin_Uploads_Content_Media_2d03e12533169.docx",
    11: "_PortalAdmin_Uploads_Content_FastAccess_3b33704e89846.docx",
    12: "_PortalAdmin_Uploads_Content_FastAccess_fcd53e9e44261.docx",
}

# Taksonomi kararı UYGULANDI (RENAME, dokumanlar/08).
_GRUP_TAKMA_ADLAR = {
    "Kamu ve Özel Hiz. Sek. ile Diğer": "Kamu ve Özel Hizmetler",
    "Kamu ve Özel Hizmetler Sektörü ile Diğer": "Kamu ve Özel Hizmetler",
    "Tarımsal Faaliyetler": "Tarımsal",
    "Ticarethane": "Kamu ve Özel Hizmetler",
    "Tarımsal Sulama": "Tarımsal",
}
_ATLA_ETIKETLERI = {
    "Genel Toplam",
    "GENEL TOPLAM",
    "Toplam",
    "İl Toplam",
    "Türkiye",
    "TÜRKİYE",
    "Pay",
    "Pay(%)",
    "Pay (%)",
    "Pay\n(%)",
    "Payı",
    "Payı (%)",
}


def grup_esle_zorunlu(metin: str) -> str | None:
    temiz = metin.strip()
    if temiz in _ATLA_ETIKETLERI:
        return None
    if temiz in _GRUP_TAKMA_ADLAR:
        return _GRUP_TAKMA_ADLAR[temiz]
    grup = _excel_grup_esle(temiz)
    if grup is None:
        raise ValueError(
            f"Tanınmayan tüketici grubu etiketi: {temiz!r} — "
            "worker/scripts/word_2019.py: _GRUP_TAKMA_ADLAR'a eklenmeli mi "
            "kontrol et (yeni bir ay yeni bir kısaltma/varyant getirmiş olabilir)."
        )
    return grup


# dokumanlar/08 devamı — 2018/2019'un T4 tablosunda "Güneş" iki AYRI kolona
# bölünmüş: "Güneş (Fotovoltaik)" (worker/parser.py:kaynak_esle() zaten
# tanıyor) ve "Güneş (Yoğunlş.)" (TANIMIYOR — "Yoğunlaştırılmış" kısaltması,
# burada AYNI kanonik "Güneş"e eklendi). t4_oku() bu ikisini TOPLAR, ayrı
# satır üretmez (bkz. modül notu). **Ay ay hücre içi boşluk/satır sonu
# tutarsız** (örn. Ekim'de "Güneş \n(Yoğunlş.)" — parantezden önce satır
# sonu var, diğer aylarda yok) — bu yüzden `kaynak_esle_zorunlu()` eşleme
# ÖNCESİ tüm iç boşluk/satır sonlarını TEK boşluğa indirgiyor.
_KAYNAK_TAKMA_ADLAR: dict[str, str] = {
    "Güneş (Yoğunlş.)": "Güneş",
}
_KAYNAK_ATLA_ETIKETLERI = {"Genel Toplam", "Toplam", "İl Toplam"}
TUM_IL_KODLARI = set(range(1, 82))


def kaynak_esle_zorunlu(metin: str) -> str | None:
    # İç boşluk/satır sonlarını normalize et (bkz. yukarıdaki not) — dış
    # boşluk zaten .strip() ile atılır.
    temiz = " ".join(metin.split())
    if temiz in _KAYNAK_ATLA_ETIKETLERI:
        return None
    if temiz in _KAYNAK_TAKMA_ADLAR:
        return _KAYNAK_TAKMA_ADLAR[temiz]
    eslesme = kaynak_esle(temiz)
    if eslesme is None:
        raise ValueError(
            f"Tanınmayan kaynak türü etiketi: {temiz!r} — "
            "worker/scripts/word_2019.py: _KAYNAK_TAKMA_ADLAR'a eklenmeli mi kontrol et."
        )
    return eslesme[0]


_IL_ADI_DUZELT: dict[str, str] = {}


def _il_adi_temizle(il_adi_ham: str) -> str:
    # Ekim 2019'da bazı il adları ORTADAN satır sonu ile bölünmüş (örn.
    # "DÜZC\nE") — yalnız satır-sonu/tab karakterleri kaldırılır (81 il
    # adının hiçbiri gerçekte bunları İÇERMEZ), GERÇEK boşluklar (örn.
    # t4_oku()'nun aradığı "Genel Toplam") KORUNUR.
    temiz = re.sub(r"[\n\r\t]", "", il_adi_ham).rstrip("* ").strip()
    return _IL_ADI_DUZELT.get(temiz, temiz)


_BILINEN_ETIKET_HATALARI: dict[tuple[int, int, str], tuple[str, int]] = {}


def _ay_yil_dogrula_kapak(
    kapak_baslik: str, beklenen_ay_adi: str, beklenen_yil: int, ay: int
) -> None:
    """word_2020.py:_ay_yil_dogrula_kapak() ile BİREBİR AYNI mantık."""
    m = re.search(r"(20\d\d)\s+Yılı\s+(\w+)\s+Ayı", kapak_baslik)
    if not m:
        raise ValueError(f"Kapak başlığından ay/yıl çıkarılamadı: {kapak_baslik!r}")
    bulunan_yil, bulunan_ay = int(m.group(1)), m.group(2)
    if bulunan_ay == beklenen_ay_adi and bulunan_yil == beklenen_yil:
        return
    bilinen = _BILINEN_ETIKET_HATALARI.get((ay, beklenen_yil, "KAPAK"))
    if bilinen == (beklenen_ay_adi, beklenen_yil):
        print(
            f"  [BİLİNEN KAYNAK HATASI] Kapak: {bulunan_ay} {bulunan_yil} "
            f"diyor ama beklenen {beklenen_ay_adi} {beklenen_yil} — devam ediliyor."
        )
        return
    raise ValueError(
        f"MANIFEST_2019 uyuşmazlığı! beklenen={beklenen_ay_adi} {beklenen_yil}, "
        f"kapağın kendi başlığı={bulunan_ay} {bulunan_yil} (başlık: {kapak_baslik!r})"
    )


def t11_oku(tbl, tarih_id: int) -> pd.DataFrame:
    """T11-karşılığı: wide format. Karar 2: Sanayi HARİÇ. word_2020.py:
    t11_oku() ile BİREBİR AYNI mantık."""
    baslik_satir = [c.text.strip() for c in tbl.rows[0].cells]
    grup_kolonlari_tam = grup_kolonlarini_coz(baslik_satir, grup_esle_zorunlu)
    grup_kolonlari = [(idx, g) for idx, g in grup_kolonlari_tam if g != "Sanayi"]
    if not grup_kolonlari:
        raise ValueError(
            f"T11: hiç grup kolonu bulunamadı (Sanayi hariç), başlık={baslik_satir}"
        )

    satirlar = []
    for row in tbl.rows[1:]:
        hucreler = [c.text.strip() for c in row.cells]
        il_adi_ham = _il_adi_temizle(hucreler[0])
        if not il_adi_ham or il_adi_ham.upper() in ("GENEL TOPLAM", "TOPLAM"):
            continue
        il_kodu = il_kodu_bul(il_adi_ham)
        if il_kodu is None:
            raise ValueError(f"T11: il_kodu bulunamadı: {il_adi_ham!r}")
        for kolon_idx, grup in grup_kolonlari:
            satirlar.append(
                {
                    "il_kodu": il_kodu,
                    "tarih_id": tarih_id,
                    "grup": grup,
                    "baglanti": "dagitim",
                    "tuketim_mwh": parse_sayi(hucreler[kolon_idx]),
                }
            )
    df = pd.DataFrame(
        satirlar, columns=["il_kodu", "tarih_id", "grup", "baglanti", "tuketim_mwh"]
    )
    beklenen = 81 * len(grup_kolonlari)
    if len(df) != beklenen:
        raise ValueError(
            f"T11: beklenen satır {beklenen} (81 il × {len(grup_kolonlari)} grup), "
            f"gerçek {len(df)}"
        )
    return df


def t10_oku(tbl, tarih_id: int, hedef_ay_yil: str) -> pd.DataFrame:
    """T10-karşılığı: uzun format. word_2020.py'nin DİNAMİK başlık-satırı
    bulma mantığıyla AYNI."""
    baslik_satir_idx = None
    for idx, row in enumerate(tbl.rows):
        if any("Tüketici Türü" in c.text for c in row.cells):
            baslik_satir_idx = idx
            break
    if baslik_satir_idx is None:
        raise ValueError("T10: 'Tüketici Türü' içeren başlık satırı bulunamadı")
    baslik_satiri = [c.text.strip() for c in tbl.rows[baslik_satir_idx].cells]
    donem_satirlari = [
        [c.text.strip() for c in tbl.rows[i].cells] for i in range(baslik_satir_idx)
    ]
    donem_satiri = [
        " ".join(satir[kolon] for satir in donem_satirlari)
        for kolon in range(len(baslik_satiri))
    ]
    baslik_metni = " ".join(baslik_satiri)
    if "Sayı" in baslik_metni:
        baslik_iceren = "Sayı"
    elif "Miktar" in baslik_metni:
        baslik_iceren = "Miktar"
    else:
        raise ValueError(f"T10: kolon başlığı tanınmadı: {baslik_satiri}")
    hedef_kolon = hedef_donem_kolonu_bul(
        donem_satiri, baslik_satiri, hedef_ay_yil, baslik_iceren
    )

    satirlar = []
    son_il_adi = ""
    for row in tbl.rows[baslik_satir_idx + 1 :]:
        hucreler = [c.text.strip() for c in row.cells]
        il_adi_ham = _il_adi_temizle(hucreler[0])
        if il_adi_ham:
            son_il_adi = il_adi_ham
        else:
            il_adi_ham = son_il_adi
        tur_ham = hucreler[1]
        if not il_adi_ham or not tur_ham:
            continue
        grup = grup_esle_zorunlu(tur_ham)
        if grup is None:
            continue
        if il_adi_ham.upper() in ("TÜRKİYE",):
            continue
        il_kodu = il_kodu_bul(il_adi_ham)
        if il_kodu is None:
            raise ValueError(f"T10: il_kodu bulunamadı: {il_adi_ham!r}")
        satirlar.append(
            {
                "il_kodu": il_kodu,
                "tarih_id": tarih_id,
                "grup": grup,
                "abone_sayisi": parse_sayi(hucreler[hedef_kolon]),
            }
        )
    df = pd.DataFrame(satirlar, columns=["il_kodu", "tarih_id", "grup", "abone_sayisi"])
    beklenen = 81 * 5
    if len(df) != beklenen:
        raise ValueError(f"T10: beklenen satır {beklenen}, gerçek {len(df)}")
    return df


def t4_oku(tbl, tarih_id: int) -> pd.DataFrame:
    """T4-karşılığı: il×kaynak matrisi, Lisanssız. word_2020.py:t4_oku()'dan
    FARKLI: birden fazla kolon AYNI kanonik kaynağa eşlenebiliyor (bkz.
    modül notu — "Güneş (Fotovoltaik)"+"Güneş (Yoğunlş.)") — bu yüzden il
    başına kaynak TOPLAMI biriktirilip TEK satır yazılıyor, kolon sayısı
    kadar satır DEĞİL."""
    baslik_satir = [c.text.strip() for c in tbl.rows[0].cells]
    kaynak_kolonlari: list[tuple[int, str]] = []
    toplam_kolon_idx: int | None = None
    for idx, hucre in enumerate(baslik_satir):
        if idx == 0:
            continue
        if hucre.strip() in ("Toplam", "Genel Toplam"):
            toplam_kolon_idx = idx
            continue
        kaynak = kaynak_esle_zorunlu(hucre)
        if kaynak is None:
            continue
        kaynak_kolonlari.append((idx, kaynak))
    if not kaynak_kolonlari:
        raise ValueError(f"T4: hiç kaynak kolonu bulunamadı, başlık={baslik_satir}")
    if toplam_kolon_idx is None:
        raise ValueError(f"T4: 'Toplam' kolonu bulunamadı, başlık={baslik_satir}")

    satirlar = []
    gorulen_iller: set[int] = set()
    tum_kaynaklar = sorted({kaynak for _, kaynak in kaynak_kolonlari})
    genel_toplam_deger: float | None = None
    for row in tbl.rows[1:]:
        hucreler = [c.text.strip() for c in row.cells]
        il_adi_ham = _il_adi_temizle(hucreler[0])
        if not il_adi_ham:
            continue
        if il_adi_ham.upper() in ("GENEL TOPLAM", "TOPLAM"):
            genel_toplam_deger = parse_sayi(hucreler[toplam_kolon_idx])
            continue
        il_kodu = il_kodu_bul(il_adi_ham)
        if il_kodu is None:
            raise ValueError(f"T4: il_kodu bulunamadı: {il_adi_ham!r}")
        gorulen_iller.add(il_kodu)
        il_toplam: dict[str, float] = dict.fromkeys(tum_kaynaklar, 0.0)
        for kolon_idx, kaynak in kaynak_kolonlari:
            deger = parse_sayi(hucreler[kolon_idx])
            il_toplam[kaynak] += deger if deger is not None else 0.0
        for kaynak, deger in il_toplam.items():
            satirlar.append(
                {
                    "il_kodu": il_kodu,
                    "tarih_id": tarih_id,
                    "kaynak": kaynak,
                    "lisans": "Lisanssız",
                    "kurulu_guc_mw": deger,
                }
            )

    eksik_iller = TUM_IL_KODLARI - gorulen_iller
    for il_kodu in sorted(eksik_iller):
        for kaynak in tum_kaynaklar:
            satirlar.append(
                {
                    "il_kodu": il_kodu,
                    "tarih_id": tarih_id,
                    "kaynak": kaynak,
                    "lisans": "Lisanssız",
                    "kurulu_guc_mw": 0.0,
                }
            )

    df = pd.DataFrame(
        satirlar, columns=["il_kodu", "tarih_id", "kaynak", "lisans", "kurulu_guc_mw"]
    )
    if genel_toplam_deger is None:
        raise ValueError("T4: 'Genel Toplam' satırı bulunamadı — doğrulama yapılamadı")
    hesaplanan = float(df["kurulu_guc_mw"].sum())
    fark = abs(hesaplanan - genel_toplam_deger)
    tolerans = max(0.5, abs(genel_toplam_deger) * 0.001)
    if fark > tolerans:
        raise ValueError(
            f"T4: hesaplanan toplam {hesaplanan:.2f} MW, tablonun kendi Genel "
            f"Toplam'ı {genel_toplam_deger:.2f} MW — fark {fark:.2f} MW toleransı "
            f"({tolerans:.2f}) aşıyor"
        )
    return df


def isle_ay(
    conn,
    *,
    klasor: Path,
    ay: int,
    actor_name: str,
    dry_run: bool = False,
) -> pipeline.IslemSonucu | None:
    """word_2020.py:isle_ay() ile BİREBİR AYNI akış."""
    dosya_adi = MANIFEST_2019[ay]
    yol = klasor / dosya_adi
    yil = 2019
    tarih_id = yil * 100 + ay
    ay_adi = ingest.AY_ADLARI[ay]
    ay_yil = f"{ay_adi} {yil}"
    source_period = f"{yil}-{ay:02d}"

    print(f"\n=== {ay_yil} — {dosya_adi} ===")

    if not dry_run:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT ib.batch_id, ib.status FROM ingestion_batch ib
                JOIN source_asset sa ON sa.source_asset_id = ib.source_asset_id
                WHERE sa.source_type = 'epdk_aylik_word' AND sa.source_period = %s
                  AND ib.parser_version = %s AND ib.status != 'failed'
                ORDER BY ib.batch_id
                """,
                (source_period, "word-2019-v1"),
            )
            mevcut = cur.fetchall()
        if mevcut:
            print(f"  [ATLA] {source_period} zaten işlenmiş: {mevcut}")
            return None

    icerik = yol.read_bytes()
    doc = Document(BytesIO(icerik))
    basliklar = basliklari_topla(doc)

    kapak_baslik = basliklar[0][1] if basliklar else ""
    _ay_yil_dogrula_kapak(kapak_baslik, ay_adi, yil, ay)
    print(f"  Kapak başlık: {kapak_baslik!r} (ay/yıl doğrulandı)")

    t11_tbl, t11_baslik = tek_aday_bul(
        basliklar,
        icerir=["Tüketici Türü Bazında Dağılımı"],
        icermez=["Karşılaştırılması", "Karşılaştırılmasının"],
        regex_disla=r"\b\w+-\w+\s+20\d\d\b",
        etiket="T11",
    )
    print(f"  T11 başlık: {t11_baslik!r}")

    t10_tbl, t10_baslik = tek_aday_bul(
        basliklar,
        icerir=["Tüketici Sayısının", "İl", "Karşılaştırılması"],
        icermez=["Dağıtım Bölgesi"],
        etiket="T10",
    )
    print(f"  T10 başlık: {t10_baslik!r}")

    tuketim_ham = t11_oku(t11_tbl, tarih_id)
    print(
        f"  T11: {len(tuketim_ham)} satır, grup={sorted(tuketim_ham['grup'].unique())}, "
        f"toplam={tuketim_ham['tuketim_mwh'].sum():,.2f} MWh"
    )

    abone_ham: pd.DataFrame | None
    try:
        abone_ham = t10_oku(t10_tbl, tarih_id, ay_yil)
    except ValueError as e:
        abone_ham = None
        print(
            f"  [T10 KAYNAKTA YOK] bu ay için tablo yapısal olarak il-ONLY "
            f"(grup boyutu yok) — kapsam_disi olarak işaretlenecek: {e}"
        )
    if abone_ham is not None:
        print(
            f"  T10: {len(abone_ham)} satır, grup={sorted(abone_ham['grup'].unique())}, "
            f"toplam={abone_ham['abone_sayisi'].sum():,.0f} abone"
        )

    if dry_run:
        print("  [DRY-RUN] DB'ye yazılmadı.")
        return None

    source_asset_id = ingest.kaynak_asset_olustur(
        conn,
        source_type="epdk_aylik_word",
        dosya_adi=dosya_adi,
        icerik=icerik,
        donem_tipi="aylik",
        source_period=source_period,
        uploaded_by=None,
    )
    batch_id = ingest.batch_olustur(conn, source_asset_id, "word-2019-v1", "1")
    if not ingest.batch_sahiplen(conn, batch_id):
        print(f"  [ATLA] batch_id={batch_id} zaten sahiplenilmiş/işlenmiş.")
        return None

    ingest.dim_tarih_getir_veya_olustur(conn, tarih_id)

    sonuc = pipeline.IslemSonucu(batch_id=batch_id)
    audit_tablolar: dict[str, dict[str, object]] = {}

    dogrulanan_t = kpi.dogrula_tuketim(tuketim_ham)
    yuklenen_t, atlanan_t = ingest.fact_tuketim_yukle(
        conn, dogrulanan_t.kabul, batch_id
    )
    sonuc.tablolar["fact_tuketim"] = pipeline.TabloSonucu(
        toplam=len(tuketim_ham),
        red=len(dogrulanan_t.red),
        karantina=len(dogrulanan_t.karantina),
        yuklenen=yuklenen_t,
        atlanan=atlanan_t,
    )
    audit_tablolar["fact_tuketim"] = {
        "toplam": len(tuketim_ham),
        "red": len(dogrulanan_t.red),
        "karantina": len(dogrulanan_t.karantina),
        "yuklenen": yuklenen_t,
        "atlanan": atlanan_t,
        "red_satirlari": dogrulanan_t.red.to_dict("records"),
    }

    if abone_ham is not None:
        dogrulanan_a = kpi.dogrula_abone(abone_ham)
        yuklenen_a, atlanan_a = ingest.fact_abone_yukle(
            conn, dogrulanan_a.kabul, batch_id
        )
        sonuc.tablolar["fact_abone"] = pipeline.TabloSonucu(
            toplam=len(abone_ham),
            red=len(dogrulanan_a.red),
            karantina=len(dogrulanan_a.karantina),
            yuklenen=yuklenen_a,
            atlanan=atlanan_a,
        )
        audit_tablolar["fact_abone"] = {
            "toplam": len(abone_ham),
            "red": len(dogrulanan_a.red),
            "karantina": len(dogrulanan_a.karantina),
            "yuklenen": yuklenen_a,
            "atlanan": atlanan_a,
            "red_satirlari": dogrulanan_a.red.to_dict("records"),
        }
        if dogrulanan_a.red.shape[0] or dogrulanan_a.karantina.shape[0]:
            print(
                f"  [DİKKAT] fact_abone: red={len(dogrulanan_a.red)} karantina={len(dogrulanan_a.karantina)}"
            )
    else:
        yuklenen_a = 0
        audit_tablolar["fact_abone"] = {
            "toplam": 0,
            "not": "kaynakta yok — T10 tablosu bu ay için yapısal olarak il-ONLY "
            "(grup boyutu yok), aşağıda kapsam_disi ile işaretlendi.",
        }

    if dogrulanan_t.red.shape[0] or dogrulanan_t.karantina.shape[0]:
        print(
            f"  [DİKKAT] fact_tuketim: red={len(dogrulanan_t.red)} karantina={len(dogrulanan_t.karantina)}"
        )

    toplam_satir = len(tuketim_ham) + (len(abone_ham) if abone_ham is not None else 0)
    toplam_yuklenen = yuklenen_t + yuklenen_a
    toplam_atlanan = toplam_satir - toplam_yuklenen

    ingest.batch_durumu_guncelle(
        conn,
        batch_id,
        "running",
        total_row_count=toplam_satir,
        accepted_row_count=toplam_yuklenen,
        rejected_row_count=toplam_atlanan,
    )
    ingest.audit_log_yaz(
        conn,
        table_name="ingestion_batch",
        record_id=batch_id,
        action_type="INSERT",
        actor_name=actor_name,
        payload={
            "olay": "ingest_tamamlandi",
            "tarih_id": tarih_id,
            "kaynak": "word_2019",
            "tablolar": audit_tablolar,
            "not": "T13/T1-karşılığı bu turda YOK (Karar 1 & 3, aşağıda kapsam_disi "
            "ile işaretlendi). T4-karşılığı AYRI bir batch'te — bkz. isle_ay_t4().",
        },
    )

    pipeline.kapsam_disi_isaretle(
        conn,
        tarih_id=tarih_id,
        fact_tablosu="fact_serbest_tuketici",
        sebep="Word (.docx) kaynağında Serbest Tüketici tablosu hiç bulunmuyor "
        "(dokumanlar/08_word_2016_2022_kapsam.md Bulgu 3).",
        karar_referansi="Karar 1",
    )
    if abone_ham is None:
        pipeline.kapsam_disi_isaretle(
            conn,
            tarih_id=tarih_id,
            fact_tablosu="fact_abone",
            sebep="Word (.docx) kaynağındaki Tüketici Sayısı tablosu bu ay için "
            "yapısal olarak il-ONLY (tüketici türü/grup kırılımı hiç yok) "
            "(dokumanlar/08_word_2016_2022_kapsam.md).",
            karar_referansi="08_word_2016_2022_kapsam.md",
        )

    uygun, sebep = pipeline.otomatik_onaya_uygun(sonuc)
    print(f"  otomatik_onaya_uygun() = {uygun}" + (f" ({sebep})" if sebep else ""))
    print(
        "  [NOT] onayla ÇAĞRILMADI (gece-boyu kural) — batch running/is_active=false kalıyor."
    )
    return sonuc


def isle_ay_t4(
    conn,
    *,
    klasor: Path,
    ay: int,
    actor_name: str,
    dry_run: bool = False,
) -> pipeline.IslemSonucu | None:
    """T4-karşılığı (fact_uretim, Lisanssız) için AYRı batch — word_2020.py:
    isle_ay_t4() ile BİREBİR AYNI mantık."""
    dosya_adi = MANIFEST_2019[ay]
    yol = klasor / dosya_adi
    yil = 2019
    tarih_id = yil * 100 + ay
    ay_adi = ingest.AY_ADLARI[ay]
    source_period = f"{yil}-{ay:02d}"
    parser_version = "word-2019-t4-v1"

    print(f"\n=== T4 {ay_adi} {yil} — {dosya_adi} ===")

    if not dry_run:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT ib.batch_id, ib.status FROM ingestion_batch ib
                JOIN source_asset sa ON sa.source_asset_id = ib.source_asset_id
                WHERE sa.source_type = 'epdk_aylik_word' AND sa.source_period = %s
                  AND ib.parser_version = %s AND ib.status != 'failed'
                ORDER BY ib.batch_id
                """,
                (source_period, parser_version),
            )
            mevcut = cur.fetchall()
        if mevcut:
            print(f"  [ATLA] T4 {source_period} zaten işlenmiş: {mevcut}")
            return None

    icerik = yol.read_bytes()
    doc = Document(BytesIO(icerik))
    basliklar = basliklari_topla(doc)

    t4_tbl, t4_baslik = t4_tablosunu_bul(basliklar)
    print(f"  T4 başlık: {t4_baslik!r}")

    uretim_ham = t4_oku(t4_tbl, tarih_id)
    print(
        f"  T4: {len(uretim_ham)} satır, kaynak={sorted(uretim_ham['kaynak'].unique())}, "
        f"toplam={uretim_ham['kurulu_guc_mw'].sum():,.2f} MW (Genel Toplam ile doğrulandı)"
    )

    if dry_run:
        print("  [DRY-RUN] DB'ye yazılmadı.")
        return None

    source_asset_id = ingest.kaynak_asset_olustur(
        conn,
        source_type="epdk_aylik_word",
        dosya_adi=dosya_adi,
        icerik=icerik,
        donem_tipi="aylik",
        source_period=source_period,
        uploaded_by=None,
    )
    batch_id = ingest.batch_olustur(conn, source_asset_id, parser_version, "1")
    if not ingest.batch_sahiplen(conn, batch_id):
        print(f"  [ATLA] batch_id={batch_id} zaten sahiplenilmiş/işlenmiş.")
        return None

    ingest.dim_tarih_getir_veya_olustur(conn, tarih_id)

    sonuc = pipeline.IslemSonucu(batch_id=batch_id)
    dogrulanan = kpi.dogrula_uretim(uretim_ham)
    yuklenen, atlanan = ingest.fact_uretim_yukle(conn, dogrulanan.kabul, batch_id)
    sonuc.tablolar["fact_uretim"] = pipeline.TabloSonucu(
        toplam=len(uretim_ham),
        red=len(dogrulanan.red),
        karantina=0,
        yuklenen=yuklenen,
        atlanan=atlanan,
    )
    audit_tablolar = {
        "fact_uretim": {
            "toplam": len(uretim_ham),
            "red": len(dogrulanan.red),
            "yuklenen": yuklenen,
            "atlanan": atlanan,
            "red_satirlari": dogrulanan.red.to_dict("records"),
        }
    }

    if dogrulanan.red.shape[0]:
        print(f"  [DİKKAT] fact_uretim: red={len(dogrulanan.red)}")

    ingest.batch_durumu_guncelle(
        conn,
        batch_id,
        "running",
        total_row_count=len(uretim_ham),
        accepted_row_count=yuklenen,
        rejected_row_count=len(uretim_ham) - yuklenen,
    )
    ingest.audit_log_yaz(
        conn,
        table_name="ingestion_batch",
        record_id=batch_id,
        action_type="INSERT",
        actor_name=actor_name,
        payload={
            "olay": "ingest_tamamlandi",
            "tarih_id": tarih_id,
            "kaynak": "word_2019_t4",
            "tablolar": audit_tablolar,
            "not": "Yalnız T4 (Lisanssız) - T1 (Lisanslı) kaynakta yok (Karar 3), "
            "aşağıda kapsam_disi ile işaretlendi.",
        },
    )

    pipeline.kapsam_disi_isaretle(
        conn,
        tarih_id=tarih_id,
        fact_tablosu="fact_uretim",
        nitelik="lisans_durumu=Lisanslı",
        sebep="Word (.docx) kaynağında Lisanslı kurulu güç için il×kaynak birleşik "
        "tablo yok (dokumanlar/08_word_2016_2022_kapsam.md Bulgu 2).",
        karar_referansi="Karar 3",
    )

    uygun, sebep = pipeline.otomatik_onaya_uygun(sonuc)
    print(f"  otomatik_onaya_uygun() = {uygun}" + (f" ({sebep})" if sebep else ""))
    print(
        "  [NOT] onayla ÇAĞRILMADI (gece-boyu kural) — batch running/is_active=false kalıyor."
    )
    return sonuc


def isle_ay_ulke_geneli(
    conn,
    *,
    klasor: Path,
    ay: int,
    actor_name: str,
    dry_run: bool = False,
) -> pipeline.IslemSonucu | None:
    """fact_tuketim_ulke_geneli için AYRI batch — T4'ünkiyle (isle_ay_t4)
    AYNI desen: kendi parser_version'ı ile YENİ ve BAĞIMSIZ bir batch
    zinciri, zaten onaylı fact_tuketim batch'lerine DOKUNMAZ. T11 tablosu
    BURADA YENİDEN bulunur (fact_tuketim yazımıyla hiç etkileşmez) ve
    worker/scripts/word_ortak.py:genel_toplam_satirini_oku() ile
    (t11_oku() ile AYNI grup_esle_zorunlu + parse_sayi) Genel Toplam
    satırından TÜM gruplar (Sanayi DAHİL) okunur. Karar 2 DEĞİŞMEDİ."""
    dosya_adi = MANIFEST_2019[ay]
    yol = klasor / dosya_adi
    yil = 2019
    tarih_id = yil * 100 + ay
    ay_adi = ingest.AY_ADLARI[ay]
    source_period = f"{yil}-{ay:02d}"
    parser_version = "word-2019-ulke-geneli-v1"

    print(f"\n=== Ülke Geneli {ay_adi} {yil} — {dosya_adi} ===")

    if not dry_run:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT ib.batch_id, ib.status FROM ingestion_batch ib
                JOIN source_asset sa ON sa.source_asset_id = ib.source_asset_id
                WHERE sa.source_type = 'epdk_aylik_word' AND sa.source_period = %s
                  AND ib.parser_version = %s AND ib.status != 'failed'
                ORDER BY ib.batch_id
                """,
                (source_period, parser_version),
            )
            mevcut = cur.fetchall()
        if mevcut:
            print(f"  [ATLA] Ülke Geneli {source_period} zaten işlenmiş: {mevcut}")
            return None

    icerik = yol.read_bytes()
    doc = Document(BytesIO(icerik))
    basliklar = basliklari_topla(doc)

    t11_tbl, t11_baslik = tek_aday_bul(
        basliklar,
        icerir=["Tüketici Türü Bazında Dağılımı"],
        icermez=["Karşılaştırılması", "Karşılaştırılmasının"],
        regex_disla=r"\b\w+-\w+\s+20\d\d\b",
        etiket="T11",
    )
    print(f"  T11 başlık: {t11_baslik!r}")

    degerler = genel_toplam_satirini_oku(t11_tbl, grup_esle_zorunlu)
    ulke_geneli_ham = pd.DataFrame(
        [
            {"tarih_id": tarih_id, "grup": grup, "tuketim_mwh": deger}
            for grup, deger in degerler.items()
        ],
        columns=["tarih_id", "grup", "tuketim_mwh"],
    )
    print(
        f"  Ülke Geneli: {len(ulke_geneli_ham)} satır, gruplar={sorted(degerler)}, "
        f"Sanayi={degerler.get('Sanayi')}"
    )

    if dry_run:
        print("  [DRY-RUN] DB'ye yazılmadı.")
        return None

    source_asset_id = ingest.kaynak_asset_olustur(
        conn,
        source_type="epdk_aylik_word",
        dosya_adi=dosya_adi,
        icerik=icerik,
        donem_tipi="aylik",
        source_period=source_period,
        uploaded_by=None,
    )
    batch_id = ingest.batch_olustur(conn, source_asset_id, parser_version, "1")
    if not ingest.batch_sahiplen(conn, batch_id):
        print(f"  [ATLA] batch_id={batch_id} zaten sahiplenilmiş/işlenmiş.")
        return None

    ingest.dim_tarih_getir_veya_olustur(conn, tarih_id)

    sonuc = pipeline.IslemSonucu(batch_id=batch_id)
    dogrulanan = kpi.dogrula_tuketim(ulke_geneli_ham)
    yuklenen, atlanan = ingest.fact_tuketim_ulke_geneli_yukle(
        conn, dogrulanan.kabul, batch_id
    )
    sonuc.tablolar["fact_tuketim_ulke_geneli"] = pipeline.TabloSonucu(
        toplam=len(ulke_geneli_ham),
        red=len(dogrulanan.red),
        karantina=len(dogrulanan.karantina),
        yuklenen=yuklenen,
        atlanan=atlanan,
    )
    audit_tablolar = {
        "fact_tuketim_ulke_geneli": {
            "toplam": len(ulke_geneli_ham),
            "red": len(dogrulanan.red),
            "karantina": len(dogrulanan.karantina),
            "yuklenen": yuklenen,
            "atlanan": atlanan,
            "red_satirlari": dogrulanan.red.to_dict("records"),
        }
    }
    if dogrulanan.red.shape[0] or dogrulanan.karantina.shape[0]:
        print(
            f"  [DİKKAT] fact_tuketim_ulke_geneli: red={len(dogrulanan.red)} "
            f"karantina={len(dogrulanan.karantina)}"
        )

    ingest.batch_durumu_guncelle(
        conn,
        batch_id,
        "running",
        total_row_count=len(ulke_geneli_ham),
        accepted_row_count=yuklenen,
        rejected_row_count=len(ulke_geneli_ham) - yuklenen,
    )
    ingest.audit_log_yaz(
        conn,
        table_name="ingestion_batch",
        record_id=batch_id,
        action_type="INSERT",
        actor_name=actor_name,
        payload={
            "olay": "ingest_tamamlandi",
            "tarih_id": tarih_id,
            "kaynak": "word_2019_ulke_geneli",
            "tablolar": audit_tablolar,
            "not": "T11 tablosunun kendi Genel Toplam satırından, il kırılımı "
            "olmayan ülke geneli değerler (Sanayi DAHİL) — bkz. dokumanlar/"
            "06_canli_veri_operasyon_gunlugu.md 2026-09-05 kaydı.",
        },
    )

    uygun, sebep = pipeline.otomatik_onaya_uygun(sonuc)
    print(f"  otomatik_onaya_uygun() = {uygun}" + (f" ({sebep})" if sebep else ""))
    print(
        "  [NOT] onayla ÇAĞRILMADI (gece-boyu kural) — batch running/is_active=false kalıyor."
    )
    return sonuc


def main() -> int:
    ap = argparse.ArgumentParser(
        description="EPP: 2019 Word raporlarını yükle (tek seferlik)"
    )
    ap.add_argument("--ay", type=int, choices=range(1, 13), help="Yalnız bu ayı işle")
    ap.add_argument("--klasor", type=Path, default=KLASOR_VARSAYILAN)
    ap.add_argument("--actor", default="manual-cli:word-2019")
    ap.add_argument(
        "--dry-run", action="store_true", help="Yalnız parse et, DB'ye YAZMA"
    )
    ap.add_argument("--t4", action="store_true", help="T11/T10 yerine YALNIZ T4'ü işle")
    ap.add_argument(
        "--ulke-geneli",
        action="store_true",
        help="T11/T10 yerine YALNIZ fact_tuketim_ulke_geneli'yi işle (ayrı batch)",
    )
    # KESİN KURAL: --onayla YOK, BİLİNÇLİ OLARAK.
    args = ap.parse_args()

    aylar = [args.ay] if args.ay else sorted(MANIFEST_2019)
    isleyici = (
        isle_ay_ulke_geneli
        if args.ulke_geneli
        else (isle_ay_t4 if args.t4 else isle_ay)
    )

    if args.dry_run:
        for ay in aylar:
            try:
                isleyici(
                    None, klasor=args.klasor, ay=ay, actor_name=args.actor, dry_run=True
                )
            except Exception as e:  # noqa: BLE001 - gece-boyu kural: bir ay hata verirse BEKLEMEDE say, sıradakine geç
                print(f"  [BEKLEMEDE] {ingest.AY_ADLARI[ay]} 2019 dry-run'da hata: {e}")
        return 0

    import psycopg

    database_url = get_database_url()
    if not database_url:
        print("HATA: DATABASE_URL tanımlı değil.")
        return 1

    with psycopg.connect(database_url, prepare_threshold=None) as conn:
        for ay in aylar:
            try:
                sonuc = isleyici(conn, klasor=args.klasor, ay=ay, actor_name=args.actor)
                conn.commit()
            except Exception as e:  # noqa: BLE001 - gece-boyu kural: bir ay hata verirse BEKLEMEDE say, sıradakine geç
                conn.rollback()
                print(
                    f"  [BEKLEMEDE] {ingest.AY_ADLARI[ay]} 2019 işlenirken hata oluştu, "
                    f"bu ay ATLANDI (ROLLBACK), sıradaki aya geçiliyor: {e}"
                )
                continue
            if sonuc is None:
                continue
            # KESİN KURAL: onayla.py / pipeline.batch_onayla() BURADA ÇAĞRILMAZ.
    return 0


if __name__ == "__main__":
    sys.exit(main())
