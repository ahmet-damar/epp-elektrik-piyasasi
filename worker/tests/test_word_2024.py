"""EPP — word_2024.py regresyon testleri (worker/scripts/word_2024.py).

Canlı .docx dosyalarına ya da DATABASE_URL'e bağımlı DEĞİL — synthetic
in-memory docx tabloları (python-docx `Document().add_table()`, gerçek il
adları `worker.parser._IL_ADI_KANONIK`'ten) ve saf fonksiyon testleriyle
`t11_oku`/`t10_oku`/`t4_oku`/`grup_esle_zorunlu`/`kaynak_esle_zorunlu`/
`_ay_yil_dogrula`'yı doğrudan doğrular — CI'nin 'Worker (lint · types ·
validation)' job'ında (DATABASE_URL yok) da çalışır.

2023'ün aksine 2024 tek şablon (dedike `_il_adi_temizle` yok, tek grup
alias'ı var) — bkz. dokumanlar/09_PROJE_DURUMU.md "Bilinen açık maddeler".
"""

from __future__ import annotations

import pytest
from docx import Document

from worker.parser import _IL_ADI_KANONIK
from worker.scripts.word_2024 import (
    _ay_yil_dogrula,
    grup_esle_zorunlu,
    kaynak_esle_zorunlu,
    t4_oku,
    t10_oku,
    t11_oku,
)

TUM_ILLER = [_IL_ADI_KANONIK[kod] for kod in sorted(_IL_ADI_KANONIK)]
assert len(TUM_ILLER) == 81


# ---------------------------------------------------------------------------
# Saf fonksiyon testleri — docx gerekmez
# ---------------------------------------------------------------------------


def test_grup_esle_zorunlu_kamu_ozel_kisaltma_esler() -> None:
    """Mart 2024 T10-karşılığında bulunan tek kısaltma — 2023 Eylül ile
    AYNI ('Kamu/Özel/Diğer')."""
    assert grup_esle_zorunlu("Kamu/Özel/Diğer") == "Kamu ve Özel Hizmetler"


def test_grup_esle_zorunlu_kanonik_gruplar_direkt_esler() -> None:
    assert grup_esle_zorunlu("Aydınlatma") == "Aydınlatma"
    assert grup_esle_zorunlu("Mesken") == "Mesken"
    assert grup_esle_zorunlu("Sanayi") == "Sanayi"
    assert grup_esle_zorunlu("Tarımsal Faaliyetler") == "Tarımsal"
    assert (
        grup_esle_zorunlu("Kamu ve Özel Hizmetler Sektörü ile Diğer")
        == "Kamu ve Özel Hizmetler"
    )


def test_grup_esle_zorunlu_atla_etiketleri_none_doner() -> None:
    assert grup_esle_zorunlu("Genel Toplam") is None
    assert grup_esle_zorunlu("Pay") is None


def test_grup_esle_zorunlu_bilinmeyen_etiket_hata_verir() -> None:
    with pytest.raises(ValueError, match="Tanınmayan tüketici grubu"):
        grup_esle_zorunlu("Hiç Bilinmeyen Bir Grup")


def test_kaynak_esle_zorunlu_linyit_dahil_bilinen_kaynaklar() -> None:
    """Nisan 2024'ten itibaren Linyit eklendi (5→6 kaynak sütunu,
    dokumanlar/06_canli_veri_operasyon_gunlugu.md 2026-09-02 T4 turu)."""
    assert kaynak_esle_zorunlu("Biyokütle") is not None
    assert kaynak_esle_zorunlu("Linyit") is not None
    assert kaynak_esle_zorunlu("Genel Toplam") is None


def test_ay_yil_dogrula_normal_uyusma_gecer() -> None:
    _ay_yil_dogrula("Tablo 2.6 Nisan 2024 Döneminde ...", "Nisan", 2024, "T11")


def test_ay_yil_dogrula_uyusmazlik_reddedilir() -> None:
    with pytest.raises(ValueError, match="MANIFEST_2024 uyuşmazlığı"):
        _ay_yil_dogrula("Tablo 2.6 Ocak 2024 Döneminde ...", "Şubat", 2024, "T11")


# ---------------------------------------------------------------------------
# Yapısal testler — synthetic in-memory docx tabloları
# ---------------------------------------------------------------------------


def _tablo_ekle(satirlar: list[list[str]]):  # type: ignore[no-untyped-def]
    doc = Document()
    satir_sayisi, kolon_sayisi = len(satirlar), len(satirlar[0])
    tbl = doc.add_table(rows=satir_sayisi, cols=kolon_sayisi)
    for i, satir in enumerate(satirlar):
        for j, deger in enumerate(satir):
            tbl.rows[i].cells[j].text = deger
    return tbl


def test_t11_oku_81_il_dogru_toplam_ve_sanayi_haric() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu ve Özel Hizmetler Sektörü ile Diğer",
        "Genel Toplam",
        "Pay",
    ]
    satirlar = [baslik]
    for il in TUM_ILLER:
        satirlar.append([il, "10,0", "20,0", "999,0", "5,0", "8,0", "1042,0", "1,0"])
    tbl = _tablo_ekle(satirlar)

    df = t11_oku(tbl, tarih_id=202404)

    assert len(df) == 81 * 4  # Sanayi HARİÇ (Karar 2)
    assert "Sanayi" not in set(df["grup"])
    assert (df["baglanti"] == "dagitim").all()
    assert df[df["grup"] == "Mesken"]["tuketim_mwh"].eq(20.0).all()


def test_t10_oku_hedef_donem_kolonu_metin_aramasiyla_bulunur() -> None:
    donem_satiri = ["", "", "2023\nMart", "2023\nMart", "2024\nMart", "2024\nMart", ""]
    baslik_satiri = [
        "İl Adı",
        "Tüketici Türü",
        "Sayı",
        "Pay(%)",
        "Sayı",
        "Pay(%)",
        "Değişim (%)",
    ]
    satirlar = [donem_satiri, baslik_satiri]
    gruplar = [
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu/Özel/Diğer",
    ]
    for il in TUM_ILLER:
        for grup in gruplar:
            satirlar.append([il, grup, "100", "1,0", "130", "1,0", "30,0"])
    tbl = _tablo_ekle(satirlar)

    df = t10_oku(tbl, tarih_id=202403, hedef_ay_yil="2024 Mart")

    assert len(df) == 81 * 5
    assert df["il_kodu"].nunique() == 81
    assert (df["abone_sayisi"] == 130).all()


def test_t4_oku_eksik_il_acikca_sifirlanir_ve_genel_toplam_dogrulanir() -> None:
    baslik = [
        "İLLER",
        "Biyokütle",
        "Doğal Gaz",
        "Güneş",
        "Hidrolik",
        "Rüzgar",
        "Linyit",
        "Toplam",
    ]
    satirlar = [baslik]
    satirlar.append(["Eskişehir", "1,0", "2,0", "3,0", "4,0", "5,0", "6,0", "21,0"])
    satirlar.append(["Ankara", "0,0", "0,0", "10,0", "0,0", "0,0", "0,0", "10,0"])
    satirlar.append(["Genel Toplam", "1,0", "2,0", "13,0", "4,0", "5,0", "6,0", "31,0"])
    tbl = _tablo_ekle(satirlar)

    df = t4_oku(tbl, tarih_id=202404)

    assert df["il_kodu"].nunique() == 81
    assert float(df["kurulu_guc_mw"].sum()) == pytest.approx(31.0)
    mugla_kodu = next(kod for kod, ad in _IL_ADI_KANONIK.items() if ad == "Muğla")
    assert df[df["il_kodu"] == mugla_kodu]["kurulu_guc_mw"].eq(0.0).all()


def test_t4_oku_genel_toplam_uyusmazliginda_hata_verir() -> None:
    baslik = ["İLLER", "Biyokütle", "Toplam"]
    satirlar = [
        baslik,
        ["Eskişehir", "5,0", "5,0"],
        ["Genel Toplam", "5,0", "999,0"],
    ]
    tbl = _tablo_ekle(satirlar)

    with pytest.raises(ValueError, match="Genel"):
        t4_oku(tbl, tarih_id=202404)
