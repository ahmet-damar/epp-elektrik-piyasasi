"""EPP — word_2025.py regresyon testleri (worker/scripts/word_2025.py).

Canlı .docx dosyalarına ya da DATABASE_URL'e bağımlı DEĞİL — synthetic
in-memory docx tabloları (python-docx `Document().add_table()`, gerçek il
adları `worker.parser._IL_ADI_KANONIK`'ten) ve saf fonksiyon testleriyle
`t11_oku`/`t10_oku`/`t4_oku`/`grup_esle_zorunlu`/`kaynak_esle_zorunlu`/
`_il_adi_temizle`/`_ay_yil_dogrula`'yı doğrudan doğrular — CI'nin
'Worker (lint · types · validation)' job'ında (DATABASE_URL yok) da çalışır.

2025 tek şablon (12 ayın tamamı ön-taramada kontrol edildi) — tek grup
alias'ı (2024 Mart ile aynı); `_il_adi_temizle` 2025'te fiilen tetiklenmedi
ama ucuz bir önlem olarak duruyor (bkz. modül notu). Bu dosya
dokumanlar/09_PROJE_DURUMU.md'nin "Bilinen açık maddeler" listesindeki
son maddeyi (word_2023/2024/2025 regresyon testi) kapatır.
"""

from __future__ import annotations

import pytest
from docx import Document

from worker.parser import _IL_ADI_KANONIK
from worker.scripts.word_2025 import (
    _ay_yil_dogrula,
    _il_adi_temizle,
    grup_esle_zorunlu,
    kaynak_esle_zorunlu,
    t4_oku,
    t10_oku,
    t11_oku,
)

TUM_ILLER = [_IL_ADI_KANONIK[kod] for kod in sorted(_IL_ADI_KANONIK)]
assert len(TUM_ILLER) == 81


# ---------------------------------------------------------------------------
# Saf fonksiyon testleri — docx gerekmez
# ---------------------------------------------------------------------------


def test_grup_esle_zorunlu_kamu_ozel_kisaltma_esler() -> None:
    assert grup_esle_zorunlu("Kamu/Özel/Diğer") == "Kamu ve Özel Hizmetler"


def test_grup_esle_zorunlu_kanonik_gruplar_direkt_esler() -> None:
    assert grup_esle_zorunlu("Aydınlatma") == "Aydınlatma"
    assert grup_esle_zorunlu("Mesken") == "Mesken"
    assert grup_esle_zorunlu("Sanayi") == "Sanayi"
    assert grup_esle_zorunlu("Tarımsal Faaliyetler") == "Tarımsal"
    assert (
        grup_esle_zorunlu("Kamu ve Özel Hizmetler Sektörü ile Diğer")
        == "Kamu ve Özel Hizmetler"
    )


def test_grup_esle_zorunlu_atla_etiketleri_none_doner() -> None:
    assert grup_esle_zorunlu("Genel Toplam") is None
    assert grup_esle_zorunlu("Pay") is None


def test_grup_esle_zorunlu_bilinmeyen_etiket_hata_verir() -> None:
    with pytest.raises(ValueError, match="Tanınmayan tüketici grubu"):
        grup_esle_zorunlu("Hiç Bilinmeyen Bir Grup")


def test_kaynak_esle_zorunlu_bilinen_kaynaklar() -> None:
    assert kaynak_esle_zorunlu("Biyokütle") is not None
    assert kaynak_esle_zorunlu("Linyit") is not None
    assert kaynak_esle_zorunlu("Genel Toplam") is None


def test_il_adi_temizle_ucuz_onlem_normal_isimleri_bozmaz() -> None:
    """2025'te fiilen tetiklenmedi ama fonksiyon 2023'ün dipnot/inceltme
    düzeltmelerini hâlâ uyguluyor — normal il adlarını bozmadığı
    doğrulanmalı."""
    assert _il_adi_temizle("İstanbul") == "İstanbul"
    assert _il_adi_temizle("Adıyaman* ") == "Adıyaman"
    assert _il_adi_temizle("HAKKÂRİ") == "HAKKARİ"


def test_ay_yil_dogrula_normal_uyusma_gecer() -> None:
    _ay_yil_dogrula("Tablo 2.6 Nisan 2025 Döneminde ...", "Nisan", 2025, "T11")


def test_ay_yil_dogrula_uyusmazlik_reddedilir() -> None:
    with pytest.raises(ValueError, match="MANIFEST_2025 uyuşmazlığı"):
        _ay_yil_dogrula("Tablo 2.6 Ocak 2025 Döneminde ...", "Şubat", 2025, "T11")


# ---------------------------------------------------------------------------
# Yapısal testler — synthetic in-memory docx tabloları
# ---------------------------------------------------------------------------


def _tablo_ekle(satirlar: list[list[str]]):  # type: ignore[no-untyped-def]
    doc = Document()
    satir_sayisi, kolon_sayisi = len(satirlar), len(satirlar[0])
    tbl = doc.add_table(rows=satir_sayisi, cols=kolon_sayisi)
    for i, satir in enumerate(satirlar):
        for j, deger in enumerate(satir):
            tbl.rows[i].cells[j].text = deger
    return tbl


def test_t11_oku_81_il_dogru_toplam_ve_sanayi_haric() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu ve Özel Hizmetler Sektörü ile Diğer",
        "Genel Toplam",
        "Pay",
    ]
    satirlar = [baslik]
    for il in TUM_ILLER:
        satirlar.append([il, "10,0", "20,0", "999,0", "5,0", "8,0", "1042,0", "1,0"])
    tbl = _tablo_ekle(satirlar)

    df = t11_oku(tbl, tarih_id=202504)

    assert len(df) == 81 * 4  # Sanayi HARİÇ (Karar 2)
    assert "Sanayi" not in set(df["grup"])
    assert (df["baglanti"] == "dagitim").all()
    assert df[df["grup"] == "Mesken"]["tuketim_mwh"].eq(20.0).all()


def test_t10_oku_hedef_donem_kolonu_metin_aramasiyla_bulunur() -> None:
    """Nisan 2025'te bulunan 5-illik toplu 'Aydınlatma' kırmızı-satır
    deseni (bkz. dokumanlar/06_canli_veri_operasyon_gunlugu.md,
    2026-09-02 devam) t10_oku'nun kendisini etkilemiyor — yalnız
    kırmızı-satır incelemesi/aktivasyon aşamasında ele alınıyor, bu
    testin kapsamı dışında."""
    donem_satiri = [
        "",
        "",
        "2024\nNisan",
        "2024\nNisan",
        "2025\nNisan",
        "2025\nNisan",
        "",
    ]
    baslik_satiri = [
        "İl Adı",
        "Tüketici Türü",
        "Sayı",
        "Pay(%)",
        "Sayı",
        "Pay(%)",
        "Değişim (%)",
    ]
    satirlar = [donem_satiri, baslik_satiri]
    gruplar = [
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu/Özel/Diğer",
    ]
    for il in TUM_ILLER:
        for grup in gruplar:
            satirlar.append([il, grup, "100", "1,0", "140", "1,0", "40,0"])
    tbl = _tablo_ekle(satirlar)

    df = t10_oku(tbl, tarih_id=202504, hedef_ay_yil="2025 Nisan")

    assert len(df) == 81 * 5
    assert df["il_kodu"].nunique() == 81
    assert (df["abone_sayisi"] == 140).all()


def test_t4_oku_eksik_il_acikca_sifirlanir_ve_genel_toplam_dogrulanir() -> None:
    baslik = [
        "İLLER",
        "Biyokütle",
        "Doğal Gaz",
        "Güneş",
        "Hidrolik",
        "Rüzgar",
        "Linyit",
        "Toplam",
    ]
    satirlar = [baslik]
    satirlar.append(["Eskişehir", "1,0", "2,0", "3,0", "4,0", "5,0", "6,0", "21,0"])
    satirlar.append(["Ankara", "0,0", "0,0", "10,0", "0,0", "0,0", "0,0", "10,0"])
    satirlar.append(["Genel Toplam", "1,0", "2,0", "13,0", "4,0", "5,0", "6,0", "31,0"])
    tbl = _tablo_ekle(satirlar)

    df = t4_oku(tbl, tarih_id=202504)

    assert df["il_kodu"].nunique() == 81
    assert float(df["kurulu_guc_mw"].sum()) == pytest.approx(31.0)
    mugla_kodu = next(kod for kod, ad in _IL_ADI_KANONIK.items() if ad == "Muğla")
    assert df[df["il_kodu"] == mugla_kodu]["kurulu_guc_mw"].eq(0.0).all()


def test_t4_oku_genel_toplam_uyusmazliginda_hata_verir() -> None:
    baslik = ["İLLER", "Biyokütle", "Toplam"]
    satirlar = [
        baslik,
        ["Eskişehir", "5,0", "5,0"],
        ["Genel Toplam", "5,0", "999,0"],
    ]
    tbl = _tablo_ekle(satirlar)

    with pytest.raises(ValueError, match="Genel"):
        t4_oku(tbl, tarih_id=202504)
