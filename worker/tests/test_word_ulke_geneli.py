"""EPP — worker/scripts/word_ortak.py:grup_kolonlarini_coz() ve
genel_toplam_satirini_oku() testleri (fact_tuketim_ulke_geneli, 2026-09-05).

Canlı .docx dosyalarına ya da DATABASE_URL'e bağımlı DEĞİL — synthetic
in-memory docx tabloları (python-docx `Document().add_table()`) ile.
CI'nin 'Worker (lint · types · validation)' job'ında (DATABASE_URL yok)
da çalışır.

Kullanıcı talebiyle (2026-09-05) eklenen 4 şart burada karşılanıyor:
1. genel_toplam_satirini_oku() kendi eşleme mantığını ÇIKARMIYOR —
   grup_kolonlarini_coz() + gerçek yıl modüllerinin (word_2016/word_2024)
   grup_esle_zorunlu()'sunu kullanıyor (test_genel_toplam_..._word_2024_
   grup_esle_zorunlu_ile_calisir, test_2016_tipi_...).
2. "Genel Toplam" satırı il-filtreleme döngüsünden BAĞIMSIZ, ham
   tbl.rows üzerinden aranıyor — test_genel_toplam_satiri_il_listesinden_
   filtrelenmis_olsa_bile_bulunur bunu KANITLAR (araya il_kodu_bul()'un
   reddedeceği bozuk "il" satırları serpiştirilmiş, yine de doğru okunur).
3. (mutabakat sorgusu testi worker/ katmanında değil, backfill script'inde
   — bkz. dokumanlar/06_canli_veri_operasyon_gunlugu.md).
4. test_genel_toplam_2024_mart_bilinen_dogru_sayi — GERÇEK 2024-03
   docx'inden (Tablo 2.6 Mart 2024) doğrudan okunup doğrulanmış Genel
   Toplam satırı (Sanayi=9.103.299,20 MWh) sabit kodlanmış, sentetik
   tabloya YERLEŞTİRİLİP fonksiyonun bu bilinen doğru sayıyı ürettiği
   test ediliyor — yalnız yapısal değil, SAYISAL bir referansa kilitli.
"""

from __future__ import annotations

import pandas as pd
import pytest
from docx import Document

from worker import kpi
from worker.scripts.word_2016 import grup_esle_zorunlu as grup_esle_zorunlu_2016
from worker.scripts.word_2024 import grup_esle_zorunlu as grup_esle_zorunlu_2024
from worker.scripts.word_ortak import genel_toplam_satirini_oku, grup_kolonlarini_coz


def _tablo_ekle(satirlar: list[list[str]]):  # type: ignore[no-untyped-def]
    doc = Document()
    satir_sayisi, kolon_sayisi = len(satirlar), len(satirlar[0])
    tbl = doc.add_table(rows=satir_sayisi, cols=kolon_sayisi)
    for i, satir in enumerate(satirlar):
        for j, deger in enumerate(satir):
            tbl.rows[i].cells[j].text = deger
    return tbl


# ---------------------------------------------------------------------------
# grup_kolonlarini_coz — Sanayi DAHİL tüm gruplar (t11_oku'nun filtrelediği
# listeden FARKLI olarak)
# ---------------------------------------------------------------------------


def test_grup_kolonlarini_coz_sanayi_dahil_tum_gruplari_dondurur() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu ve Özel Hizmetler Sektörü ile Diğer",
        "Genel Toplam",
        "Pay",
    ]
    kolonlar = grup_kolonlarini_coz(baslik, grup_esle_zorunlu_2024)
    gruplar = {g for _, g in kolonlar}
    assert gruplar == {
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal",
        "Kamu ve Özel Hizmetler",
    }
    assert "Genel Toplam" not in gruplar and "Pay" not in gruplar


def test_grup_kolonlarini_coz_hic_grup_kolonu_yoksa_hata_verir() -> None:
    with pytest.raises(ValueError, match="Hiç grup kolonu bulunamadı"):
        grup_kolonlarini_coz(["İller", "Genel Toplam", "Pay"], grup_esle_zorunlu_2024)


# ---------------------------------------------------------------------------
# Senaryo (a) — 2016 tipi: tek, basit tablo (il kırılımı yok, 2016'nın
# kendi eski taksonomisi: Tarımsal Sulama/Ticarethane)
# ---------------------------------------------------------------------------


def test_genel_toplam_2016_tipi_tek_tablo() -> None:
    """2016'nın T11 (il×grup, GENİŞ format) tablosu — kolon başlıkları
    2016'nın ESKİ taksonomisi (Tarımsal Sulama/Ticarethane, RENAME
    öncesi). genel_toplam_satirini_oku() HER ZAMAN bu GENİŞ tabloyu
    (fact_tuketim için zaten bulunan AYNI Table nesnesini) okur — ayrı,
    dar formatlı bir "Tablo 2.3" ARANMAZ (bkz. modül notu)."""
    baslik = [
        "İl",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Sulama",
        "Ticarethane",
        "Genel Toplam",
    ]
    satirlar = [
        baslik,
        ["Ankara", "1,0", "1,0", "1,0", "1,0", "1,0", "5,0"],
        [
            "Genel Toplam",
            "457.141,72",
            "5.324.982,80",
            "6.717.607,62",
            "150.366,52",
            "5.283.331,39",
            "17.933.430,06",
        ],
    ]
    tbl = _tablo_ekle(satirlar)
    degerler = genel_toplam_satirini_oku(tbl, grup_esle_zorunlu_2016)

    assert degerler["Sanayi"] == pytest.approx(6717607.62)
    assert degerler["Aydınlatma"] == pytest.approx(457141.72)
    # 2016'nın taksonomi RENAME'i (word_2016.py:_GRUP_TAKMA_ADLAR) burada
    # da aynen geçerli — "Tarımsal Sulama" -> "Tarımsal", "Ticarethane" ->
    # "Kamu ve Özel Hizmetler" (t11_oku ile AYNI eşleme, iki farklı mantık
    # OLUŞMADIĞININ kanıtı).
    assert degerler["Tarımsal"] == pytest.approx(150366.52)
    assert degerler["Kamu ve Özel Hizmetler"] == pytest.approx(5283331.39)
    assert set(degerler) == {
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal",
        "Kamu ve Özel Hizmetler",
    }


# ---------------------------------------------------------------------------
# Senaryo (b) — GERÇEK 2024-03 verisiyle, bilinen doğru sayıya kilitli
# (kullanıcı talebi madde 4). Kaynak: MANIFEST_2024[3]'ün "Tablo 2.6 Mart
# 2024 Döneminde ... İl ve Tüketici Türü Bazında Dağılımı" tablosunun
# GERÇEK başlık + Genel Toplam satırı (2026-09-05'te doğrudan docx'ten
# okunup doğrulandı — dokumanlar/06_canli_veri_operasyon_gunlugu.md).
# ---------------------------------------------------------------------------


def test_genel_toplam_2024_mart_bilinen_dogru_sayi() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Kamu ve Özel Hizmetler Sektörü ile Diğer",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Genel Toplam",
        "Pay",
    ]
    # Yalnız Genel Toplam satırı gerçek — il satırları test amaçlı kısa
    # tutuldu (genel_toplam_satirini_oku il satırlarına HİÇ bakmaz, bkz.
    # test_genel_toplam_satiri_il_listesinden_filtrelenmis_olsa_bile_bulunur).
    satirlar = [
        baslik,
        ["Ankara", "1,0", "1,0", "1,0", "1,0", "1,0", "5,0", "0,01%"],
        [
            "Genel Toplam",
            "479.970,65",
            "5.428.802,37",
            "6.040.499,24",
            "9.103.299,20",
            "254.230,83",
            "21.306.802,29",
            "100,00%",
        ],
    ]
    tbl = _tablo_ekle(satirlar)
    degerler = genel_toplam_satirini_oku(tbl, grup_esle_zorunlu_2024)

    assert degerler["Sanayi"] == pytest.approx(9103299.20)
    assert degerler["Aydınlatma"] == pytest.approx(479970.65)
    assert degerler["Kamu ve Özel Hizmetler"] == pytest.approx(5428802.37)
    assert degerler["Mesken"] == pytest.approx(6040499.24)
    assert degerler["Tarımsal"] == pytest.approx(254230.83)
    # kendi kendine tutarlılık: 5 grubun toplamı tablonun kendi "Genel
    # Toplam" (toplam kolonu, burada okunmuyor ama satırın kendi metni
    # ile) değeriyle örtüşüyor mü diye ayrıca (worker/kpi.py'siz) doğrula.
    assert sum(degerler.values()) == pytest.approx(21306802.29, rel=1e-6)


# ---------------------------------------------------------------------------
# Senaryo (c) — "Genel Toplam" satırı yoksa ValueError (tahmin YOK)
# ---------------------------------------------------------------------------


def test_genel_toplam_satiri_yoksa_hata_verir() -> None:
    baslik = ["İl", "Aydınlatma", "Sanayi"]
    satirlar = [
        baslik,
        ["Ankara", "457.141,72", "6.717.607,62"],
        ["İstanbul", "1,0", "2,0"],
    ]
    tbl = _tablo_ekle(satirlar)
    with pytest.raises(ValueError, match="Genel Toplam.*bulunamadı"):
        genel_toplam_satirini_oku(tbl, grup_esle_zorunlu_2016)


# ---------------------------------------------------------------------------
# Kullanıcı talebi madde 2 — "Genel Toplam" satırı, il-eşlemesi sırasında
# filtrelenmiş/reddedilmiş olsa bile HAM tbl.rows'tan doğru bulunur.
# ---------------------------------------------------------------------------


def test_genel_toplam_satiri_il_listesinden_filtrelenmis_olsa_bile_bulunur() -> None:
    """Araya, t11_oku()'nun il-döngüsünün il_kodu_bul() ile REDDEDECEĞİ
    (bozuk/tanınmayan) sahte 'il' satırları serpiştirildi — bu fonksiyon
    o döngüye hiç girmediğinden, bu satırlar sorun ÇIKARMADAN Genel
    Toplam satırına ulaşılabildiğini kanıtlar."""
    satirlar = [
        ["İller", "Aydınlatma", "Sanayi", "Genel Toplam", "Pay"],
        ["BÖYLE_BİR_İL_YOK", "1,0", "2,0", "3,0", "1,0%"],
        ["İLLER", "0,0", "0,0", "0,0", "0,0%"],  # 2017 tekrar-başlık anomalisi
        ["", "", "", "", ""],  # boş satır
        ["Genel Toplam", "457.141,72", "6.717.607,62", "7.174.749,34", "100,00%"],
    ]
    tbl = _tablo_ekle(satirlar)
    degerler = genel_toplam_satirini_oku(tbl, grup_esle_zorunlu_2016)

    assert degerler["Sanayi"] == pytest.approx(6717607.62)
    assert degerler["Aydınlatma"] == pytest.approx(457141.72)


# ---------------------------------------------------------------------------
# Negatif değer — genel_toplam_satirini_oku() KENDİSİ reddetmez (ham
# okuma), ama kpi.dogrula_tuketim() (isle_ay_ulke_geneli()'nde AYNEN
# fact_tuketim ile kullanılan doğrulayıcı) reddeder — mevcut mutabakat
# kurallarıyla tutarlı.
# ---------------------------------------------------------------------------


def test_negatif_deger_ham_okumada_gecer_ama_dogrula_tuketim_reddeder() -> None:
    baslik = ["İl", "Aydınlatma", "Sanayi"]
    satirlar = [
        baslik,
        ["Ankara", "1,0", "1,0"],
        ["Genel Toplam", "457.141,72", "-7.270,79"],
    ]
    tbl = _tablo_ekle(satirlar)
    degerler = genel_toplam_satirini_oku(tbl, grup_esle_zorunlu_2016)
    assert degerler["Sanayi"] == pytest.approx(-7270.79)

    df = pd.DataFrame(
        [
            {"tarih_id": 202612, "grup": grup, "tuketim_mwh": deger}
            for grup, deger in degerler.items()
        ]
    )
    dogrulanan = kpi.dogrula_tuketim(df)
    assert len(dogrulanan.red) == 1
    assert dogrulanan.red.iloc[0]["grup"] == "Sanayi"
    assert len(dogrulanan.kabul) == 1
    assert dogrulanan.kabul.iloc[0]["grup"] == "Aydınlatma"
