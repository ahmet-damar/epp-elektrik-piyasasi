"""EPP — word_2023.py regresyon testleri (worker/scripts/word_2023.py).

Canlı .docx dosyalarına ya da DATABASE_URL'e bağımlı DEĞİL — synthetic
in-memory docx tabloları (python-docx `Document().add_table()`, gerçek il
adları `worker.parser._IL_ADI_KANONIK`'ten) ve saf fonksiyon testleriyle
`t11_oku`/`t10_oku`/`t4_oku`/`grup_esle_zorunlu`/`kaynak_esle_zorunlu`/
`_il_adi_temizle`/`_ay_yil_dogrula`'yı doğrudan doğrular — CI'nin
'Worker (lint · types · validation)' job'ında (DATABASE_URL yok) da çalışır.

Bkz. dokumanlar/09_PROJE_DURUMU.md "Bilinen açık maddeler" — bu dosya
word_2023.py için o eksiği kapatır (test_word_2016..2022.py zaten vardı).
"""

from __future__ import annotations

import pytest
from docx import Document

from worker.parser import _IL_ADI_KANONIK
from worker.scripts.word_2023 import (
    _ay_yil_dogrula,
    _il_adi_temizle,
    grup_esle_zorunlu,
    kaynak_esle_zorunlu,
    t4_oku,
    t10_oku,
    t11_oku,
)

TUM_ILLER = [_IL_ADI_KANONIK[kod] for kod in sorted(_IL_ADI_KANONIK)]
assert len(TUM_ILLER) == 81


# ---------------------------------------------------------------------------
# Saf fonksiyon testleri — docx gerekmez
# ---------------------------------------------------------------------------


def test_grup_esle_zorunlu_kamu_ozel_dort_varyant_esler() -> None:
    """2023'ün 'Kamu ve Özel Hizmetler' için 4 farklı kısaltma varyantı
    kullandığı bulundu (Ocak/Şubat T11 birbirinden bile farklı, Eylül/Nisan
    T10 2024 ile örtüşüyor) — hepsi AYNI kanonik gruba eşlenmeli."""
    assert (
        grup_esle_zorunlu("Kamu ve Özel Hiz. Sek. ile Diğer")
        == "Kamu ve Özel Hizmetler"
    )
    assert grup_esle_zorunlu("Kamu/Özel Hiz. Sek./Diğer") == "Kamu ve Özel Hizmetler"
    assert grup_esle_zorunlu("Kamu/Özel/Diğer") == "Kamu ve Özel Hizmetler"
    assert grup_esle_zorunlu("Kamu/Özel/ Diğer") == "Kamu ve Özel Hizmetler"


def test_grup_esle_zorunlu_kanonik_gruplar_direkt_esler() -> None:
    assert grup_esle_zorunlu("Aydınlatma") == "Aydınlatma"
    assert grup_esle_zorunlu("Mesken") == "Mesken"
    assert grup_esle_zorunlu("Sanayi") == "Sanayi"
    assert grup_esle_zorunlu("Tarımsal Faaliyetler") == "Tarımsal"


def test_grup_esle_zorunlu_atla_etiketleri_none_doner() -> None:
    assert grup_esle_zorunlu("Genel Toplam") is None
    assert grup_esle_zorunlu("Pay") is None


def test_grup_esle_zorunlu_bilinmeyen_etiket_hata_verir() -> None:
    with pytest.raises(ValueError, match="Tanınmayan tüketici grubu"):
        grup_esle_zorunlu("Hiç Bilinmeyen Bir Grup")


def test_kaynak_esle_zorunlu_bilinen_kaynaklar() -> None:
    assert kaynak_esle_zorunlu("Biyokütle") is not None
    assert kaynak_esle_zorunlu("Doğal Gaz") is not None
    assert kaynak_esle_zorunlu("Genel Toplam") is None


def test_il_adi_temizle_adiyaman_dipnot_yildizi_kaldirilir() -> None:
    """Ocak 2023 T11'de 'ADIYAMAN*' — 2023 depremiyle ilgili bir EPDK
    dipnot yıldızı, il adının bir parçası DEĞİL."""
    assert _il_adi_temizle("ADIYAMAN*") == "ADIYAMAN"
    assert _il_adi_temizle("Adıyaman* ") == "Adıyaman"


def test_il_adi_temizle_hakkari_inceltme_isareti_duzeltilir() -> None:
    """Ağustos 2023'te eski yazım 'HAKKÂRİ' (^) kullanılmış — standart
    yazım 'HAKKARİ'ye çevrilmeli, worker/parser.py DEĞİŞMEDEN."""
    assert _il_adi_temizle("HAKKÂRİ") == "HAKKARİ"


def test_ay_yil_dogrula_normal_uyusma_gecer() -> None:
    _ay_yil_dogrula("Tablo 2.6 Mart 2023 Döneminde ...", "Mart", 2023, "T11")


def test_ay_yil_dogrula_uyusmazlik_reddedilir() -> None:
    with pytest.raises(ValueError, match="MANIFEST_2023 uyuşmazlığı"):
        _ay_yil_dogrula("Tablo 2.6 Ocak 2023 Döneminde ...", "Şubat", 2023, "T11")


# ---------------------------------------------------------------------------
# Yapısal testler — synthetic in-memory docx tabloları
# ---------------------------------------------------------------------------


def _tablo_ekle(satirlar: list[list[str]]):  # type: ignore[no-untyped-def]
    doc = Document()
    satir_sayisi, kolon_sayisi = len(satirlar), len(satirlar[0])
    tbl = doc.add_table(rows=satir_sayisi, cols=kolon_sayisi)
    for i, satir in enumerate(satirlar):
        for j, deger in enumerate(satir):
            tbl.rows[i].cells[j].text = deger
    return tbl


def test_t11_oku_81_il_dogru_toplam_ve_sanayi_haric() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu ve Özel Hiz. Sek. ile Diğer",
        "Genel Toplam",
        "Pay",
    ]
    satirlar = [baslik]
    for il in TUM_ILLER:
        satirlar.append([il, "10,0", "20,0", "999,0", "5,0", "8,0", "1042,0", "1,0"])
    tbl = _tablo_ekle(satirlar)

    df = t11_oku(tbl, tarih_id=202303)

    assert len(df) == 81 * 4  # Sanayi HARİÇ (Karar 2)
    assert "Sanayi" not in set(df["grup"])
    assert (df["baglanti"] == "dagitim").all()
    assert df[df["grup"] == "Mesken"]["tuketim_mwh"].eq(20.0).all()


def test_t10_oku_hedef_donem_kolonu_metin_aramasiyla_bulunur() -> None:
    """2023'ün T10-karşılığı, 2022'nin il×grup 'Dönemler Arası
    Karşılaştırma' yapısıyla AYNI (sabit rows[0]=dönem/rows[1]=başlık) —
    hedef dönem kolonu 'Sayı' iceren başlıkla metin aramasıyla bulunuyor."""
    donem_satiri = [
        "",
        "",
        "2022\nEylül",
        "2022\nEylül",
        "2023\nEylül",
        "2023\nEylül",
        "",
    ]
    baslik_satiri = [
        "İl Adı",
        "Tüketici Türü",
        "Sayı",
        "Pay(%)",
        "Sayı",
        "Pay(%)",
        "Değişim (%)",
    ]
    satirlar = [donem_satiri, baslik_satiri]
    gruplar = [
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu/Özel/Diğer",
    ]
    for il in TUM_ILLER:
        for grup in gruplar:
            satirlar.append([il, grup, "100", "1,0", "120", "1,0", "20,0"])
    tbl = _tablo_ekle(satirlar)

    df = t10_oku(tbl, tarih_id=202309, hedef_ay_yil="2023 Eylül")

    assert len(df) == 81 * 5
    assert df["il_kodu"].nunique() == 81
    # Hedef "2023 Eylül" kolonu (120) seçilmeli, önceki yıl (100) DEĞİL
    assert (df["abone_sayisi"] == 120).all()


def test_t4_oku_eksik_il_acikca_sifirlanir_ve_genel_toplam_dogrulanir() -> None:
    baslik = [
        "İLLER",
        "Biyokütle",
        "Doğal Gaz",
        "Güneş",
        "Hidrolik",
        "Rüzgar",
        "Toplam",
    ]
    satirlar = [baslik]
    satirlar.append(["Eskişehir", "1,0", "2,0", "3,0", "4,0", "5,0", "15,0"])
    satirlar.append(["Ankara", "0,0", "0,0", "10,0", "0,0", "0,0", "10,0"])
    satirlar.append(["Genel Toplam", "1,0", "2,0", "13,0", "4,0", "5,0", "25,0"])
    tbl = _tablo_ekle(satirlar)

    df = t4_oku(tbl, tarih_id=202303)

    assert df["il_kodu"].nunique() == 81  # eksik iller AÇIKÇA 0 ile tamamlanmış
    assert float(df["kurulu_guc_mw"].sum()) == pytest.approx(25.0)
    mugla_kodu = next(kod for kod, ad in _IL_ADI_KANONIK.items() if ad == "Muğla")
    assert df[df["il_kodu"] == mugla_kodu]["kurulu_guc_mw"].eq(0.0).all()


def test_t4_oku_genel_toplam_uyusmazliginda_hata_verir() -> None:
    baslik = ["İLLER", "Biyokütle", "Toplam"]
    satirlar = [
        baslik,
        ["Eskişehir", "5,0", "5,0"],
        ["Genel Toplam", "5,0", "999,0"],  # kasıtlı uyuşmazlık
    ]
    tbl = _tablo_ekle(satirlar)

    with pytest.raises(ValueError, match="Genel"):
        t4_oku(tbl, tarih_id=202303)
