"""EPP — word_2022.py regresyon testleri (worker/scripts/word_2022.py).

Canlı .docx dosyalarına ya da DATABASE_URL'e bağımlı DEĞİL — synthetic
in-memory docx tabloları (python-docx `Document().add_table()`, gerçek il
adları `worker.parser._IL_ADI_KANONIK`'ten) ve saf fonksiyon testleriyle
`t11_oku`/`t10_oku`/`t4_oku`/`grup_esle_zorunlu`/`kaynak_esle_zorunlu`/
`_il_adi_temizle`/`_ay_yil_dogrula`'yı doğrudan doğrular — CI'nin
'Worker (lint · types · validation)' job'ında (DATABASE_URL yok) da çalışır.

Bkz. dokumanlar/07_word_parser_kapsam.md "Açık kalanlar" madde 1 (regresyon
testi eksikliği) — bu dosya word_2022.py için bu eksiği kapatır; word_2023/
2024/2025.py için hâlâ AÇIK (bkz. dokumanlar/SABAH_OZETI.md).
"""

from __future__ import annotations

import pytest
from docx import Document

from worker.parser import _IL_ADI_KANONIK
from worker.scripts.word_2022 import (
    _ay_yil_dogrula,
    _il_adi_temizle,
    grup_esle_zorunlu,
    kaynak_esle_zorunlu,
    t4_oku,
    t10_oku,
    t11_oku,
)

TUM_ILLER = [_IL_ADI_KANONIK[kod] for kod in sorted(_IL_ADI_KANONIK)]
assert len(TUM_ILLER) == 81


# ---------------------------------------------------------------------------
# Saf fonksiyon testleri — docx gerekmez
# ---------------------------------------------------------------------------


def test_grup_esle_zorunlu_mayis_aralik_kanonik_esler() -> None:
    """Mayıs-Aralık 2022'nin gerçek etiketleri (dokumanlar/08 Bulgu 6)
    kanonik worker/parser.py GRUP_ESLEME kümesine doğru eşleşmeli."""
    assert grup_esle_zorunlu("Aydınlatma") == "Aydınlatma"
    assert grup_esle_zorunlu("Mesken") == "Mesken"
    assert grup_esle_zorunlu("Sanayi") == "Sanayi"
    assert grup_esle_zorunlu("Tarımsal Faaliyetler") == "Tarımsal"
    assert (
        grup_esle_zorunlu("Kamu ve Özel Hiz. Sek. ile Diğer")
        == "Kamu ve Özel Hizmetler"
    )
    assert (
        grup_esle_zorunlu("Kamu ve Özel Hizmetler Sektörü ile Diğer")
        == "Kamu ve Özel Hizmetler"
    )


def test_grup_esle_zorunlu_atla_etiketleri_none_doner() -> None:
    assert grup_esle_zorunlu("Genel Toplam") is None
    assert grup_esle_zorunlu("GENEL TOPLAM") is None
    assert grup_esle_zorunlu("Pay(%)") is None


def test_grup_esle_zorunlu_ocak_nisan_eski_takson_kasitli_engellenir() -> None:
    """dokumanlar/08 Bulgu 5 — 'Ticarethane'/'Tarımsal Sulama' BİLİNÇLİ
    OLARAK haritalanmadı (AÇIK KARAR, Ahmet'e sorulmalı). Bu test o
    tasarım kararının kod seviyesinde GERÇEKTEN uygulandığını doğrular —
    biri yanlışlıkla bir alias eklerse bu test KIRILIR, bilinçli bir
    hatırlatma olarak."""
    with pytest.raises(ValueError, match="KASITLI"):
        grup_esle_zorunlu("Ticarethane")
    with pytest.raises(ValueError, match="KASITLI"):
        grup_esle_zorunlu("Tarımsal Sulama")


def test_kaynak_esle_zorunlu_bilinen_kaynaklar() -> None:
    assert kaynak_esle_zorunlu("Biyokütle") is not None
    assert kaynak_esle_zorunlu("Güneş") is not None
    assert kaynak_esle_zorunlu("Genel Toplam") is None


def test_il_adi_temizle_kuthahya_kaynak_yazim_hatasi_duzeltilir() -> None:
    """dokumanlar/08 devamı — Ocak/Şubat 2022 T4 tablosunda 'Kütahya' yerine
    'Küthahya' yazılmış (fazladan 'h', gerçek kaynak yazım hatası)."""
    assert _il_adi_temizle("Küthahya") == "Kütahya"
    assert _il_adi_temizle("İstanbul") == "İstanbul"
    assert _il_adi_temizle("Adıyaman* ") == "Adıyaman"


def test_ay_yil_dogrula_normal_uyusma_gecer() -> None:
    _ay_yil_dogrula("Tablo 2.6 Mayıs 2022 Döneminde ...", "Mayıs", 2022, "T11", ay=5)


def test_ay_yil_dogrula_bilinen_etiket_hatasi_gecer() -> None:
    """dokumanlar/08 Bulgu 6 — Nisan 2022 dosyasının T11 tablosu kendi
    başlığında 'Mart 2022' diyor (kaynağın kendi hatası, veri duplikasyonu
    DEĞİL) — _BILINEN_ETIKET_HATALARI bu TEK istisnayı belgeliyor."""
    _ay_yil_dogrula("Tablo 2.6 Mart 2022 Döneminde ...", "Nisan", 2022, "T11", ay=4)


def test_ay_yil_dogrula_bilinmeyen_uyusmazlik_reddedilir() -> None:
    """Bilinen istisna LİSTESİNDE OLMAYAN bir uyuşmazlık sessizce
    geçilmemeli — gerçek bir manifest hatası olabilir."""
    with pytest.raises(ValueError, match="MANIFEST_2022 uyuşmazlığı"):
        _ay_yil_dogrula("Tablo 2.6 Ocak 2022 Döneminde ...", "Şubat", 2022, "T11", ay=2)


# ---------------------------------------------------------------------------
# Yapısal testler — synthetic in-memory docx tabloları
# ---------------------------------------------------------------------------


def _tablo_ekle(satirlar: list[list[str]]):  # type: ignore[no-untyped-def]
    """Verilen hücre metinleriyle yeni bir docx tablosu oluşturur, tabloyu
    (python-docx Table nesnesi) döndürür — t11_oku/t10_oku/t4_oku bu
    nesneyi (gerçek bir .docx dosyasından geldiği gibi) doğrudan alır."""
    doc = Document()
    satir_sayisi, kolon_sayisi = len(satirlar), len(satirlar[0])
    tbl = doc.add_table(rows=satir_sayisi, cols=kolon_sayisi)
    for i, satir in enumerate(satirlar):
        for j, deger in enumerate(satir):
            tbl.rows[i].cells[j].text = deger
    return tbl


def test_t11_oku_81_il_dogru_toplam_ve_sanayi_haric() -> None:
    baslik = [
        "İller",
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal",
        "Kamu ve Özel Hizmetler",
        "Genel Toplam",
        "Pay",
    ]
    satirlar = [baslik]
    for il in TUM_ILLER:
        satirlar.append([il, "10,0", "20,0", "999,0", "5,0", "8,0", "1042,0", "1,0"])
    tbl = _tablo_ekle(satirlar)

    df = t11_oku(tbl, tarih_id=202205)

    assert len(df) == 81 * 4  # Sanayi HARİÇ (Karar 2) -> 4 grup kaldı
    assert "Sanayi" not in set(df["grup"])
    assert set(df["grup"]) == {
        "Aydınlatma",
        "Mesken",
        "Tarımsal",
        "Kamu ve Özel Hizmetler",
    }
    assert (df["baglanti"] == "dagitim").all()
    # Her il için Mesken 20,0 olmalı (parse_sayi TR biçimini doğru okumalı)
    assert df[df["grup"] == "Mesken"]["tuketim_mwh"].eq(20.0).all()


def test_t10_oku_ileri_doldurma_bos_il_hucresi() -> None:
    """dokumanlar/08 Bulgu 6 devamı — bazı aylarda (Ekim gibi) İl Adı
    hücresi yalnız ilin İLK grup satırında dolu, sonrakiler BOŞ (Word'ün
    kendi hücre-birleştirmesi) — t10_oku bir önceki dolu İl Adı'nı
    kullanmalı (05_kaynak_dosya_sozlesmesi.md T13'teki AYNI 'ileri
    doldurma' ilkesi)."""
    satirlar = [
        ["", "", "2021\nEkim", "2021\nEkim", "2022\nEkim", "2022\nEkim", ""],
        [
            "İl Adı",
            "Tüketici Türü",
            "Tüketici Sayısı",
            "Pay(%)",
            "Tüketici Sayısı",
            "Pay(%)",
            "Değişim (%)",
        ],
    ]
    gruplar = [
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal Faaliyetler",
        "Kamu ve Özel Hizmetler Sektörü ile Diğer",
    ]
    for il in TUM_ILLER:
        for i, grup in enumerate(gruplar):
            il_hucresi = il if i == 0 else ""  # yalnız ilk satırda İl Adı dolu
            satirlar.append([il_hucresi, grup, "100", "1,0", "110", "1,0", "10,0"])
    tbl = _tablo_ekle(satirlar)

    df = t10_oku(tbl, tarih_id=202210, hedef_ay_yil="Ekim 2022")

    assert len(df) == 81 * 5
    assert df["il_kodu"].nunique() == 81
    assert set(df["grup"]) == {
        "Aydınlatma",
        "Mesken",
        "Sanayi",
        "Tarımsal",
        "Kamu ve Özel Hizmetler",
    }
    # Hedef dönem "2022 Ekim" kolonu seçilmiş olmalı (110), önceki yıl (100) DEĞİL
    assert (df["abone_sayisi"] == 110).all()


def test_t4_oku_eksik_il_acikca_sifirlanir_ve_genel_toplam_dogrulanir() -> None:
    baslik = [
        "İLLER",
        "Biyokütle",
        "Doğal Gaz",
        "Güneş",
        "Hidrolik",
        "Rüzgar",
        "Toplam",
    ]
    satirlar = [baslik]
    # Yalnız 3 il satır olarak var (Artvin/Hakkari gibi bazı iller 2023'te
    # sessizce atlanmıştı, bkz. Bulgu 5 madde 4 — aynı desen burada da test edilir)
    satirlar.append(["Eskişehir", "1,0", "2,0", "3,0", "4,0", "5,0", "15,0"])
    satirlar.append(["Ankara", "0,0", "0,0", "10,0", "0,0", "0,0", "10,0"])
    satirlar.append(["İstanbul", "0,0", "0,0", "0,0", "0,0", "20,0", "20,0"])
    satirlar.append(["Genel Toplam", "1,0", "2,0", "13,0", "4,0", "25,0", "45,0"])
    tbl = _tablo_ekle(satirlar)

    df = t4_oku(tbl, tarih_id=202205)

    assert df["il_kodu"].nunique() == 81  # eksik 78 il AÇIKÇA 0 ile tamamlanmış
    assert len(df) == 81 * 5
    assert float(df["kurulu_guc_mw"].sum()) == pytest.approx(45.0)
    # Görülmeyen bir il (örn. Muğla) tüm kaynaklarda gerçekten 0.0 olmalı
    mugla_kodu = next(kod for kod, ad in _IL_ADI_KANONIK.items() if ad == "Muğla")
    assert df[df["il_kodu"] == mugla_kodu]["kurulu_guc_mw"].eq(0.0).all()


def test_t4_oku_genel_toplam_uyusmazliginda_hata_verir() -> None:
    """Genel Toplam satırı hesaplanan toplamla UYUŞMUYORSA (hizalama/parse
    sorunu şüphesi) ValueError — sessizce yanlış veri kabul edilmez."""
    baslik = ["İLLER", "Biyokütle", "Toplam"]
    satirlar = [
        baslik,
        ["Eskişehir", "5,0", "5,0"],
        ["Genel Toplam", "5,0", "999,0"],  # kasıtlı uyuşmazlık
    ]
    tbl = _tablo_ekle(satirlar)

    with pytest.raises(ValueError, match="Genel"):
        t4_oku(tbl, tarih_id=202205)
